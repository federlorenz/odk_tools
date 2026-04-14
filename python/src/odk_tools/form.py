# %% #@ Libraries imports

import pandas as pd
import copy
import xml.etree.ElementTree as ET
import docx as dcx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from re import findall
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

# %%


class ProcessQuestionnaire():
    def __init__(self):
        self.attachments = None
        self.survey = None
        self.choices = None
        self.settings = None
        self.form_title = None
        self.form_version = None
        self.languages = None

    @classmethod
    def strip_double_column(cls, df):
        df.columns = [i.replace("::", ":") for i in df.columns]
        return df

    def get_data_from_files(self, form_filename, attachements_list_filenames):

        self.survey = pd.read_excel(form_filename, na_values=[
            ' ', ''], keep_default_na=False, sheet_name="survey").dropna(how='all')
        self.choices = pd.read_excel(form_filename, sheet_name="choices", na_values=[
            ' ', ''], keep_default_na=False).dropna(how='all')
        self.settings = pd.read_excel(form_filename, sheet_name="settings", na_values=[
            ' ', ''], keep_default_na=False).dropna(how='all')
        self.form_title = self.settings['form_title'].iloc[0]
        self.form_version = self.settings['version'].iloc[0]
        attachments = {}
        for j in attachements_list_filenames:
            attachments[j] = pd.read_csv(j)
        self.attachments = attachments

    def get_data_from_odk_object(self, odk_object):
        self.survey = odk_object.survey
        self.choices = odk_object.choices
        self.settings = odk_object.settings
        self.form_title = self.settings['form_title'].iloc[0]
        self.form_version = self.settings['version'].iloc[0] if odk_object.form_is_published(
        ) else "FORM_NOT_PUBLISHED"
        self.attachments = odk_object.attachments
        self.form_is_published = odk_object.form_is_published()

    def get_languages(self):
        language = []
        for column in self.survey.columns:
            if column[:5] == "label":
                if len(column) == 5:
                    language.append("")
                else:
                    language.append(column[5:])
        self.languages = sorted(list(set(language)))

    def process(self, highlight_color={"begin_group": "4F81BD", "end_group": "B8CCE4", "begin_repeat": "9BBB59", "end_repeat": "D6E3BC", "calculate": "D9D9D9", "header_row": "919191"}, language=None, paragraph_spacing_points=3, compress_long_choices=True, to_memory_filename=False):

        document = dcx.Document()
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = dcx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = dcx.shared.Cm(0.5)
        section.bottom_margin = dcx.shared.Cm(0.5)
        section.left_margin = dcx.shared.Cm(1)
        section.right_margin = dcx.shared.Cm(1)
        heading = document.add_heading(
            f"Form title = {self.form_title}\nForm version = {str(self.form_version)}{"\n" if (language == None or language == "") else "\nForm language = "+language.split(":")[-1]}\n\n")
        heading.alignment = 1

        p = document.add_heading('Headings explained', level=2)

        p = document.add_paragraph()
        run = p.add_run('Question code')
        run.bold = True
        run = p.add_run(': The code used to identify the question')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Question type')
        run.bold = True
        run = p.add_run(': The question type. Common types are "text", "integer", "select_one", "select_multiple", "date", "time", "image, Note". Questions of type "calculate" are internal variables used in ODK to process calculations and should be ignored for the purpose of reviewing the questionnaire.')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Question')
        run.bold = True
        run = p.add_run(
            ': The question label in the specified language (default language English)')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Hint')
        run.bold = True
        run = p.add_run(
            ': The hint in the specified language (default language English)')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Select options')
        run.bold = True
        run = p.add_run(
            ': For questions of type select_one or select_multiple, these are the options that can be selected.')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Logic')
        run.bold = True
        run = p.add_run(
            ': These are the logics defined for the question. Different logic can be defined. ')
        run.italic = True
        run = p.add_run('Relevant')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(' defines the logic for showing the question or not. ')
        run.italic = True
        run = p.add_run('Default value')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(' is the value that the question assumes by default. ')
        run.italic = True
        run = p.add_run('Constrain value')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies the limits imposed on the values that can be typed in. ')
        run.italic = True
        run = p.add_run('Calculation')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies the calculations for questions of type calculate. ')
        run.italic = True
        run = p.add_run('Choice filter')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies if the selectable options for questions of type select_one or select_multiple should be shown according to some logic. ')
        run.italic = True
        run = p.add_run('Repeat count')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies the number of times that a repeat block is repeated. ')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        p = document.add_heading('Rows highlighting explained', level=2)

        legend = document.add_table(rows=0, cols=2)
        for k, v in {"A block of questions begins": "4F81BD",
                     "A block of questions ends": "B8CCE4",
                     "A repeated block of questions begins": "9BBB59",
                     "A repeated block fo questions ends": "D6E3BC",
                     "A questions of type \"calculate\"": "D9D9D9",
                     "The table headers": "919191",
                     "Cells modified after review": "FFFF00"}.items():
            row_cells = legend.add_row().cells
            paragraph = row_cells[1].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(k)
            row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = row_cells[0].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
            shading_elm = parse_xml(f'<w:shd {nsdecls('w')} w:fill="{v}"/>')
            legend.rows[-1].cells[0]._tc.get_or_add_tcPr().append(shading_elm)

        for cell in legend.columns[0].cells:
            cell.width = dcx.shared.Cm(2)
        for row in legend.rows:
            row.height = dcx.shared.Cm(1)

        p = document.add_paragraph()
        run = p.add_run('\n')
        run = p.add_run('\n')
        run.italic = True

        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        title_cells = table.rows[0].cells
        headings = ["Question code", "Question type",
                    "Question", "Hint", "Select Options", "Logic"]
        for i in range(len(headings)):
            paragraph = title_cells[i].paragraphs[0]
            run = paragraph.add_run(headings[i])
            run.bold = True

        def set_repeat_table_header(row):
            """ set repeat table row on every new page
            """
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)
            return row

        def get_choices(choice):
            choice_names = list(
                self.choices[f"label{"" if (language == None or language == "") else language}"].loc[self.choices["list_name"] == choice].map(str))
            choice_labels = list(
                self.choices["name"].loc[self.choices["list_name"] == choice].map(str))
            zipped = list(zip(choice_labels, choice_names))
            zipped = [" = ".join(x) for x in zipped]
            if compress_long_choices:
                if len(zipped) > 50:
                    zipped1 = zipped[0:25]
                    zipped2 = zipped[-25:-1]
                    zipped = zipped1 + \
                        ["...The list is longer than 50, some elements are omitted..."]+zipped2
            zipped = "\n".join(zipped)
            return zipped

        def get_from_file(file, label="label"):
            choice_list = list(self.attachments[file][label].map(str))
            if compress_long_choices:
                if len(choice_list) > 50:
                    choice_list1 = choice_list[0:25]
                    choice_list2 = choice_list[-25:-1]
                    choice_list = choice_list1 + \
                        ["...The list is longer than 50, some elements are omitted..."]+choice_list2
            return "\n".join(choice_list)

        def get_alternative_label_from_file(x):
            if type(x) is float:
                return "label"
            y = x.split(",")
            for i in y:
                if "label" in i:
                    return i.split("=")[1]
            return "label"

        def process_enclosing_variables(data):
            x = findall(r"\${.+?}", data)
            for i in x:
                data = data.replace(i, i[2:-1])
            return data

        def process_current_input(data):
            x = findall(r"\.", data)
            for i in x:
                data = data.replace(i, "input ")
            return data

        def process_string_only(column, index, cell_index):
            if not pd.isna(s[column].iloc[index]):

                paragraph = row_cells[cell_index].paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(paragraph_spacing_points)
                paragraph_format.space_after = Pt(paragraph_spacing_points)
                run = paragraph.add_run(process_enclosing_variables(
                    str(s[column].iloc[index])))
                if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

        def process_string_only_combined(cells, index):
            columns_names = ["Relevant: ", "Default value: ", "Constrain value: ",
                             "Calculation: ", "Choice filter: ", "Repeat count: "]
            column_labels = ["relevant", "default", "constraint",
                             "calculation", "choice_filter", "repeat_count"]
            paragraph = cells[5].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(paragraph_spacing_points)
            paragraph_format.space_after = Pt(paragraph_spacing_points)
            count = 0
            for j in range(len(column_labels)):
                if column_labels[j] in self.survey.columns:
                    if (not pd.isna(self.survey[column_labels[j]].iloc[index])):
                        count += 1
            for j in range(len(column_labels)):
                if column_labels[j] in self.survey.columns:
                    if not pd.isna(self.survey[column_labels[j]].iloc[index]):
                        run = paragraph.add_run(columns_names[j])
                        run.italics = True
                        run.font.color.rgb = RGBColor(50, 0, 255)
                        if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                            run.font.size = Pt(8)
                        if column_labels[j] == "constraint":
                            run = paragraph.add_run(
                                f"{process_current_input(process_enclosing_variables(str(self.survey[column_labels[j]].iloc[index])))}")
                            if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                                run.font.size = Pt(8)
                        else:
                            run = paragraph.add_run(
                                f"{process_enclosing_variables(str(self.survey[column_labels[j]].iloc[index]))}")
                            if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                                run.font.size = Pt(8)
                        count -= 1
                        if count != 0:
                            run = paragraph.add_run("\n")

        def cell_shading(index=None, index_counter=None, header_row=False):
            if index != None:
                sss = self.survey["type"].iloc[index].split(" ")[0]
                if sss in highlight_color.keys():
                    for j in range(6):
                        shading_elm = parse_xml(
                            f'<w:shd {nsdecls('w')} w:fill="{highlight_color[sss]}"/>')
                        table.rows[index+1 - index_counter].cells[j]._tc.get_or_add_tcPr().append(
                            shading_elm)
            if header_row:
                for j in range(6):
                    shading_elm = parse_xml(
                        f'<w:shd {nsdecls('w')} w:fill="{highlight_color['header_row']}"/>')
                    table.rows[0].cells[j]._tc.get_or_add_tcPr().append(
                        shading_elm)

        cell_shading(header_row=True)
        set_repeat_table_header(table.rows[0])
        index_counter = 0
        for i in range(len(self.survey)):
            s = self.survey
            c = self.choices
            if self.survey["relevant"].iloc[i] == "false()":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "start":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "end":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "deviceid":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "phonenumber":
                index_counter += 1
            else:
                row_cells = table.add_row().cells
                cell_shading(index=i, index_counter=index_counter)
                paragraph = row_cells[0].paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(paragraph_spacing_points)
                paragraph_format.space_after = Pt(paragraph_spacing_points)
                run = paragraph.add_run(s["name"].iloc[i] if not pd.isna(
                    s["name"].iloc[i]) else "")
                if self.survey["type"].iloc[i].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

                paragraph = row_cells[1].paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(paragraph_spacing_points)
                paragraph_format.space_after = Pt(paragraph_spacing_points)
                run = paragraph.add_run(s["type"].iloc[i].split(" ")[0])
                if self.survey["type"].iloc[i].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

                process_string_only(
                    f"label{"" if (language == None or language == "") else language}", i, 2)
                process_string_only(
                    f"hint{"" if (language == None or language == "") else language}", i, 3)

                if (s["type"].iloc[i].split(" ")[0] == "select_one") or (s["type"].iloc[i].split(" ")[0] == "select_multiple"):
                    row_cells[4].text = get_choices(
                        s["type"].iloc[i].split(" ")[1])
                elif (s["type"].iloc[i].split(" ")[0] == "select_one_from_file") or (s["type"].iloc[i].split(" ")[0] == "select_multiple_from_file"):
                    if type(s["parameters"].iloc[i]) is str:
                        label = get_alternative_label_from_file(
                            s["parameters"].iloc[i])
                    else:
                        label = "label"
                    row_cells[4].text = get_from_file(
                        s["type"].iloc[i].split(" ")[1], label=label)
                else:
                    row_cells[4].text = ""
                process_string_only_combined(row_cells, i)
        if to_memory_filename != False:
            out = copy.deepcopy(to_memory_filename)
            document.save(out)
            return out
        else:
            document.save(
                f"{self.form_title}{"" if (self.languages == None) else ("" if (language == None or language == "") else "-" + language.split(":")[-1])}-{"Version_"+str(self.form_version)}.docx")


# class ProcessSubmission():
#     def __init__(self, sub=None, survey=None, choices=None):
#         self.sub = sub
#         self.survey = survey
#         self.choices = choices

#     @classmethod
#     def strip_double_column(cls, df):
#         df.columns = [i.replace("::", ":") for i in df.columns]
#         return df

#     def get_data_from_odk_object(self, odk_object):
#         self.survey = Process_questionnaire.strip_double_column(
#             odk_object.survey)
#         self.choices = Process_questionnaire.strip_double_column(
#             odk_object.choices)

#     def get_languages(self):
#         language = []
#         for column in self.survey.columns:
#             if column[:5] == "label":
#                 if len(column) == 5:
#                     language.append("")
#                 else:
#                     language.append(column[5:])
#         self.languages = sorted(list(set(language)))

#     def set_cell_border(self, cell: _Cell, **kwargs):
#         """
#         Set cell`s border
#         Usage:

#         set_cell_border(
#             cell,
#             top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
#             bottom={"sz": 12, "color": "#00FF00", "val": "single"},
#             start={"sz": 24, "val": "dashed", "shadow": "true"},
#             end={"sz": 12, "val": "dashed"},
#         )
#         """
#         tc = cell._tc
#         tcPr = tc.get_or_add_tcPr()

#         # check for tag existnace, if none found, then create one
#         tcBorders = tcPr.first_child_found_in("w:tcBorders")
#         if tcBorders is None:
#             tcBorders = OxmlElement('w:tcBorders')
#             tcPr.append(tcBorders)

#         # list over all available tags
#         for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
#             edge_data = kwargs.get(edge)
#             if edge_data:
#                 tag = 'w:{}'.format(edge)

#                 # check for tag existnace, if none found, then create one
#                 element = tcBorders.find(qn(tag))
#                 if element is None:
#                     element = OxmlElement(tag)
#                     tcBorders.append(element)

#                 # looks like order of attributes is important
#                 for key in ["sz", "val", "color", "space", "shadow"]:
#                     if key in edge_data:
#                         element.set(qn('w:{}'.format(key)),
#                                     str(edge_data[key]))

#     def process(self, highlight_color={"begin_group": "4F81BD", "end_group": "B8CCE4", "begin_repeat": "9BBB59", "end_repeat": "D6E3BC", "calculate": "D9D9D9", "header_row": "919191"}, language=None, paragraph_spacing_points=3, compress_long_choices=True, to_memory_filename=False):

#         document = dcx.Document()
#         section = document.sections[-1]
#         new_width, new_height = section.page_height, section.page_width
#         section.orientation = dcx.enum.section.WD_ORIENT.PORTRAIT
#         section.page_width = new_width
#         section.page_height = new_height
#         section.top_margin = dcx.shared.Cm(0.5)
#         section.bottom_margin = dcx.shared.Cm(0.5)
#         section.left_margin = dcx.shared.Cm(1)
#         section.right_margin = dcx.shared.Cm(1)
#         heading = document.add_heading(
#             f"heading")
#         heading.alignment = 1

#         p = document.add_paragraph()
#         run = p.add_run('\n')
#         run = p.add_run('\n')
#         run.italic = True

#         table = document.add_table(rows=1, cols=3)
#         table.style = "Table Grid"
#         title_cells = table.rows[0].cells

#         headings = ["Question type", "Question",
#                     "Answer"]

#         for i in range(len(headings)):
#             paragraph = title_cells[i].paragraphs[0]
#             run = paragraph.add_run(headings[i])
#             run.bold = True

#         def set_repeat_table_header(row):
#             """ set repeat table row on every new page
#             """
#             tr = row._tr
#             trPr = tr.get_or_add_trPr()
#             tblHeader = OxmlElement('w:tblHeader')
#             tblHeader.set(qn('w:val'), "true")
#             trPr.append(tblHeader)
#             return row

#         def get_choices(choice):
#             choice_names = list(
#                 self.choices[f"label{"" if (language == None or language == "") else language}"].loc[self.choices["list_name"] == choice].map(str))
#             choice_labels = list(
#                 self.choices["name"].loc[self.choices["list_name"] == choice].map(str))
#             zipped = list(zip(choice_labels, choice_names))
#             zipped = [" = ".join(x) for x in zipped]
#             if compress_long_choices:
#                 if len(zipped) > 50:
#                     zipped1 = zipped[0:25]
#                     zipped2 = zipped[-25:-1]
#                     zipped = zipped1 + \
#                         ["...The list is longer than 50, some elements are omitted..."]+zipped2
#             zipped = "\n".join(zipped)
#             return zipped

#         def get_from_file(file, label="label"):
#             choice_list = list(self.attachments[file][label].map(str))
#             if compress_long_choices:
#                 if len(choice_list) > 50:
#                     choice_list1 = choice_list[0:25]
#                     choice_list2 = choice_list[-25:-1]
#                     choice_list = choice_list1 + \
#                         ["...The list is longer than 50, some elements are omitted..."]+choice_list2
#             return "\n".join(choice_list)

#         def get_alternative_label_from_file(x):
#             if type(x) is float:
#                 return "label"
#             y = x.split(",")
#             for i in y:
#                 if "label" in i:
#                     return i.split("=")[1]
#             return "label"

#         def process_enclosing_variables(data):
#             x = findall(r"\${.+?}", data)
#             for i in x:
#                 data = data.replace(i, i[2:-1])
#             return data

#         def process_current_input(data):
#             x = findall(r"\.", data)
#             for i in x:
#                 data = data.replace(i, "input ")
#             return data

#         def process_string_only(column, index, cell_index):
#             if not pd.isna(s[column].iloc[index]):

#                 paragraph = row_cells[cell_index].paragraphs[0]
#                 paragraph_format = paragraph.paragraph_format
#                 paragraph_format.space_before = Pt(paragraph_spacing_points)
#                 paragraph_format.space_after = Pt(paragraph_spacing_points)
#                 run = paragraph.add_run(process_enclosing_variables(
#                     str(s[column].iloc[index])))
#                 if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
#                     run.italics = True
#                     run.font.size = Pt(8)

#         def cell_shading(index=None, index_counter=None, header_row=False):
#             if index != None:
#                 sss = self.survey["type"].iloc[index].split(" ")[0]
#                 if sss in highlight_color.keys():
#                     for j in range(6):
#                         shading_elm = parse_xml(
#                             f'<w:shd {nsdecls('w')} w:fill="{highlight_color[sss]}"/>')
#                         table.rows[index+1 - index_counter].cells[j]._tc.get_or_add_tcPr().append(
#                             shading_elm)
#             if header_row:
#                 for j in range(6):
#                     shading_elm = parse_xml(
#                         f'<w:shd {nsdecls('w')} w:fill="{highlight_color['header_row']}"/>')
#                     table.rows[0].cells[j]._tc.get_or_add_tcPr().append(
#                         shading_elm)

#         # cell_shading(header_row=True)
#         set_repeat_table_header(table.rows[0])
#         index_counter = 0
#         for i in range(len(self.survey)):
#             s = self.survey
#             c = self.choices
#             if self.survey["relevant"].iloc[i] == "false()":
#                 index_counter += 1
#             elif self.survey["type"].iloc[i].split(" ")[0] == "start":
#                 index_counter += 1
#             elif self.survey["type"].iloc[i].split(" ")[0] == "end":
#                 index_counter += 1
#             elif self.survey["type"].iloc[i].split(" ")[0] == "deviceid":
#                 index_counter += 1
#             elif self.survey["type"].iloc[i].split(" ")[0] == "phonenumber":
#                 index_counter += 1
#             else:
#                 row_cells = table.add_row().cells
#                 # cell_shading(index=i, index_counter=index_counter)

#                 paragraph = row_cells[0].paragraphs[0]
#                 paragraph_format = paragraph.paragraph_format
#                 paragraph_format.space_before = Pt(paragraph_spacing_points)
#                 paragraph_format.space_after = Pt(paragraph_spacing_points)
#                 run = paragraph.add_run(s["type"].iloc[i].split(" ")[0])
#                 if self.survey["type"].iloc[i].split(" ")[0] == "calculate":
#                     run.italics = True
#                     run.font.size = Pt(8)

#                 process_string_only(
#                     f"label{"" if (language == None or language == "") else language}", i, 1)

#                 if (s["type"].iloc[i].split(" ")[0] == "select_one") or (s["type"].iloc[i].split(" ")[0] == "select_multiple"):
#                     row_cells[2].text = get_choices(
#                         s["type"].iloc[i].split(" ")[1])
#                 elif (s["type"].iloc[i].split(" ")[0] == "select_one_from_file") or (s["type"].iloc[i].split(" ")[0] == "select_multiple_from_file"):
#                     if type(s["parameters"].iloc[i]) is str:
#                         label = get_alternative_label_from_file(
#                             s["parameters"].iloc[i])
#                     else:
#                         label = "label"
#                     row_cells[2].text = get_from_file(
#                         s["type"].iloc[i].split(" ")[1], label=label)
#                 else:
#                     row_cells[2].text = ""
#         if to_memory_filename != False:
#             out = copy.deepcopy(to_memory_filename)
#             document.save(out)
#             return out
#         else:
#             document.save(
#                 f"out.docx")
